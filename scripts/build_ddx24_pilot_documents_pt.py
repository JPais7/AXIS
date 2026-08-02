from pathlib import Path

from docx import Document

from build_ddx24_pilot_documents import (
    ROOT,
    bullets,
    configure,
    font,
    paragraph,
    table,
    title_block,
)


OUT = ROOT / "data/publication/ddx24-pilot"
PROTOCOL = OUT / "Protocolo_piloto_DDX24_Joao_Pais_Diana_Koshman.docx"
SUMMARY = OUT / "Resumo_colaboradores_piloto_DDX24_Joao_Pais_Diana_Koshman.docx"


def build_protocol():
    doc = Document()
    configure(doc, "Protocolo piloto: validação independente de DDX24 em células T CD8")
    title_block(
        doc,
        "Protocolo piloto | Versão 1.0 | 30 de julho de 2026",
        "Validação laboratorial independente da expressão de DDX24 em células T CD8 periféricas",
        "Estudo-piloto exploratório de casos e controlos na espondilite anquilosante",
    )

    doc.add_heading("Sinopse do protocolo", level=1)
    table(doc, ["Elemento", "Especificação pré-definida"], [
        ["Pergunta", "A expressão de RNA de DDX24, ao nível do participante, é inferior em células T CD8 purificadas de adultos com espondilite anquilosante comparativamente a controlos saudáveis?"],
        ["Desenho", "Estudo-piloto laboratorial, exploratório, transversal, de casos e controlos."],
        ["Amostra-alvo", "12 casos e 12 controlos analisáveis; recrutar até 15 por grupo para compensar perdas. Mínimo pretendido: 10 por grupo."],
        ["Material", "Sangue periférico fresco; isolamento de PBMC seguido de enriquecimento intacto de células T CD8 ou separação celular validada."],
        ["Ensaio", "RT-qPCR para DDX24, normalizado pela média geométrica de pelo menos dois genes de referência estáveis e previamente validados."],
        ["Resultado primário", "Expressão normalizada de DDX24 por participante (ΔCt; um ΔCt superior significa menor expressão de DDX24)."],
        ["Comparação", "Casos versus controlos saudáveis, tendo o participante humano como unidade estatística."],
        ["Estado", "Proposta para avaliação ética, laboratorial e por colaboradores; não autoriza recrutamento nem recolha de amostras."],
    ], [2300, 7060])

    doc.add_heading("1. Fundamentação", level=1)
    paragraph(doc, "O AXIS identificou expressão inferior de DDX24 em duas coortes compatíveis de células T CD8 de memória/efetoras. Uma terceira coorte independente de CD8 geral apresentou a mesma direção. A estimativa primária em duas coortes foi -0,148 log2-CPM (IC 95%: -0,272 a -0,024); a sensibilidade com três coortes foi -0,145 (IC 95%: -0,249 a -0,041). Estes resultados observacionais necessitam de falsificação prospetiva numa população independente antes de qualquer interpretação causal ou terapêutica.")

    doc.add_heading("2. Objetivos e hipóteses", level=1)
    bullets(doc, [
        "Objetivo primário: estimar a diferença de expressão normalizada de DDX24 em células T CD8 periféricas purificadas entre casos e controlos.",
        "Hipótese direcional pré-definida: a expressão de DDX24 é inferior nos casos.",
        "Objetivo secundário: estimar viabilidade técnica, rendimento de RNA, falhas do ensaio e variabilidade entre participantes.",
        "Objetivos exploratórios: estudar associações com atividade da doença, HLA-B27 e tratamento, sem as apresentar como confirmatórias.",
    ])

    doc.add_heading("3. População do estudo", level=1)
    doc.add_heading("Casos", level=2)
    bullets(doc, [
        "Adultos com idade igual ou superior a 18 anos.",
        "Espondiloartrite axial radiográfica/espondilite anquilosante confirmada por médico e critérios aceites documentados.",
        "Capacidade para prestar consentimento informado e fornecer sangue periférico.",
        "Tratamento, atividade e duração da doença documentados, sem exclusões pós-hoc baseadas no resultado.",
    ])
    doc.add_heading("Controlos saudáveis", level=2)
    bullets(doc, [
        "Adultos sem doença reumática inflamatória, doença inflamatória intestinal, psoríase ou doença autoimune ativa.",
        "Emparelhamento por frequência com os casos, por faixa etária e sexo, sempre que possível.",
        "Sem infeção aguda ou vacinação dentro da janela recente definida no protocolo aprovado pela comissão de ética.",
    ])
    doc.add_heading("Exclusões comuns", level=2)
    bullets(doc, [
        "Doença febril aguda ou infeção atual clinicamente relevante.",
        "Incapacidade para prestar consentimento informado.",
        "Volume de sangue insuficiente, falha no isolamento de PBMC ou RNA fora dos critérios de qualidade.",
        "Qualquer exclusão de segurança adicional exigida pelo centro clínico.",
    ])

    doc.add_heading("4. Dimensão da amostra", level=1)
    paragraph(doc, "Este é um piloto de viabilidade e estimativa de variância, não um estudo definitivo. Com 10 a 15 participantes por grupo, a precisão será limitada e efeitos pequenos poderão não atingir significância estatística. O alvo de 12 participantes analisáveis por grupo equilibra viabilidade e estimativa da variabilidade. Um estudo confirmatório deverá ser dimensionado usando a variância observada no piloto e um efeito biologicamente relevante, e não o resultado mais favorável do piloto.")

    doc.add_heading("5. Dados clínicos e pré-analíticos", level=1)
    table(doc, ["Categoria", "Variáveis mínimas"], [
        ["Participante", "Código do estudo, idade e sexo à nascença; sem identificadores diretos no ficheiro de análise."],
        ["Doença", "Critérios diagnósticos, HLA-B27, duração, BASDAI, ASDAS, CRP e velocidade de sedimentação, quando disponíveis."],
        ["Tratamento", "AINE, DMARD convencional, inibidores de TNF, IL-17 ou JAK, corticosteroides e alterações recentes."],
        ["Colheita", "Data/hora, jejum se controlado, atraso de processamento, volume, operador e lote."],
        ["Controlo laboratorial", "Rendimento/viabilidade de PBMC, pureza CD8, concentração/integridade de RNA, lote de transcrição reversa, placa qPCR e razões de exclusão."],
    ], [2200, 7160])

    doc.add_heading("6. Fluxo laboratorial", level=1)
    table(doc, ["Etapa", "Requisito mínimo"], [
        ["Colheita", "Usar o mesmo anticoagulante e uma janela fixa entre colheita e processamento nos dois grupos."],
        ["Isolamento de PBMC", "Aplicar o mesmo protocolo validado e registar atraso, rendimento e viabilidade."],
        ["Enriquecimento CD8", "Preferir seleção negativa intacta para reduzir ativação; alternativamente, usar uma estratégia FACS pré-definida. Registar pureza."],
        ["RNA e cDNA", "Usar a mesma estratégia de extração e quantidade de entrada; incluir controlos de RNA e transcrição reversa."],
        ["RT-qPCR", "Validar especificidade e eficiência de 90-110%; usar réplicas técnicas, controlo sem molde e controlo sem transcriptase reversa."],
        ["Genes de referência", "Avaliar candidatos como RPLP0, TBP, HPRT1 ou PPIA e selecionar pelo menos dois genes estáveis; não assumir estabilidade de GAPDH."],
        ["Lotes", "Distribuir casos e controlos entre dias, extrações e placas; aleatorizar posições dos poços."],
        ["Cegamento", "Usar amostras codificadas para que a equipa laboratorial não necessite de conhecer o grupo."],
    ], [1800, 7560])

    doc.add_heading("7. Regras de controlo de qualidade", level=1)
    bullets(doc, [
        "Definir previamente limites para viabilidade de PBMC, pureza CD8, qualidade de RNA e variação entre réplicas.",
        "Repetir uma medição apenas perante falha técnica documentada, nunca por o resultado biológico ser inesperado.",
        "Resolver réplicas discordantes através de regra escrita e conservar os valores originais.",
        "Fechar as exclusões antes de abrir os grupos e reportar cada exclusão com uma razão não identificável.",
    ])

    doc.add_heading("8. Plano de análise estatística", level=1)
    bullets(doc, [
        "Unidade estatística: participante; as réplicas técnicas são agregadas dentro do participante.",
        "Resultado primário: ΔCt calculado com a média geométrica dos genes de referência validados; reportar a diferença entre grupos e IC 95%.",
        "Comparação primária: teste de Welch bilateral ou modelo linear equivalente; a hipótese direcional não será convertida num teste unilateral.",
        "Reportar 2^-ΔΔCt como efeito secundário interpretável, acompanhado de incerteza.",
        "Executar apenas uma sensibilidade ajustada parcimoniosa com grupo, idade e sexo, se a completude e a amostra o permitirem.",
        "Mostrar todos os valores individuais e análises com e sem observações influentes justificadas.",
        "Não interpretar ausência de significância num piloto pequeno como prova de ausência de efeito.",
    ])

    doc.add_heading("9. Controlo de viés e reprodutibilidade", level=1)
    bullets(doc, [
        "Registar temporalmente o protocolo e o programa de análise antes de retirar o cegamento.",
        "Harmonizar colheita e processamento entre grupos.",
        "Conservar Ct brutos anonimizados, mapas de placa, decisões de qualidade, código, versões e somas de verificação.",
        "Separar a análise confirmatória de DDX24 de genes e vias exploratórios.",
        "Pedir a um segundo analista que verifique exclusões, grupos e resultado primário.",
    ])

    doc.add_heading("10. Ética e governação", level=1)
    paragraph(doc, "Nenhum recrutamento ou colheita pode começar antes da aprovação da comissão de ética responsável e da autorização das instituições clínica e laboratorial. O consentimento informado deverá abranger colheita, análise molecular, variáveis clínicas codificadas, conservação, análise secundária e publicação. A chave de identificação permanecerá no centro clínico, separada do conjunto de análise. As regras locais de proteção de dados e amostras biológicas terão de ser confirmadas.")

    doc.add_heading("11. Marcos de viabilidade", level=1)
    table(doc, ["Marco", "Critério para avançar"], [
        ["Parceria", "Responsáveis clínico e laboratorial nomeados e responsabilidades aceites."],
        ["Ética", "Protocolo, consentimento, formulário clínico e plano de dados aprovados."],
        ["Ensaio técnico", "Teste não comparativo confirma pureza CD8, rendimento de RNA e desempenho de RT-qPCR."],
        ["Recrutamento", "Pelo menos 10 participantes analisáveis por grupo, com processamento equilibrado."],
        ["Análise", "Base de dados fechada e cega; controlo de qualidade aplicado; efeito por participante reportado."],
        ["Decisão", "Usar variância e viabilidade para planear replicação; só avançar para perturbação funcional após replicação credível."],
    ], [2200, 7160])

    doc.add_heading("12. Responsabilidades a confirmar", level=1)
    table(doc, ["Função", "Responsabilidade"], [
        ["Investigador clínico principal", "Diagnóstico, elegibilidade, consentimento, segurança e governação clínica."],
        ["Responsável laboratorial", "Procedimentos, biossegurança, isolamento celular, validação e controlo de qualidade."],
        ["Responsável pelos dados", "Pseudonimização, proteção da chave e controlo de acessos."],
        ["Responsável estatístico", "Plano congelado, verificações e relatório reprodutível."],
        ["João Pais e Diana Koshman", "Pergunta científica, síntese da evidência, coordenação e interpretação; funções finais a acordar."],
    ], [2600, 6760])

    doc.add_heading("13. Produtos do piloto", level=1)
    bullets(doc, [
        "Protocolo e consentimento aprovados.",
        "Procedimento laboratorial e validação de RT-qPCR.",
        "Conjunto anonimizado de dados clínicos e ΔCt por participante.",
        "Resultado primário com incerteza, fluxo de qualidade e exclusões completas.",
        "Decisão fundamentada para replicar, reformular ou terminar a hipótese.",
    ])
    OUT.mkdir(parents=True, exist_ok=True)
    doc.save(PROTOCOL)


def build_summary():
    doc = Document()
    configure(doc, "Resumo para colaboradores: piloto de validação de DDX24")
    title_block(
        doc,
        "Proposta de colaboração",
        "Piloto de validação de DDX24 na espondilite anquilosante",
        "Um estudo pequeno, falsificável e centrado no participante em células T CD8 periféricas purificadas",
    )
    doc.add_heading("Porquê este estudo?", level=1)
    paragraph(doc, "Em três coortes públicas independentes de single-cell, a expressão de DDX24 ao nível do participante foi inferior na espondilite anquilosante/espondiloartrite axial. A síntese primária compatível incluiu 47 participantes; a sensibilidade em CD8 geral incluiu 51. A certeza permanece baixa devido ao pequeno número de coortes, covariáveis clínicas incompletas e uma coorte muito pequena. Um piloto laboratorial independente é o próximo teste decisivo.")
    table(doc, ["Elemento", "Proposta"], [
        ["Pergunta", "A expressão de RNA de DDX24 é inferior em células T CD8 purificadas de casos comparativamente a controlos?"],
        ["Participantes", "Alvo: 12 casos e 12 controlos; recrutar até 15/grupo; mínimo analisável: 10/grupo."],
        ["Ensaio", "Isolamento de PBMC, enriquecimento intacto de CD8, extração de RNA e RT-qPCR validado."],
        ["Resultado", "ΔCt normalizado por participante usando pelo menos dois genes de referência estáveis."],
        ["Controlo de viés", "Processamento equilibrado, amostras codificadas, critérios de qualidade pré-definidos e estatística por participante."],
        ["Interpretação", "Replicação exploratória; sem afirmação causal, diagnóstica ou terapêutica."],
    ], [2100, 7260])
    doc.add_heading("O que procuramos num colaborador", level=1)
    bullets(doc, [
        "Acesso clínico a participantes bem caracterizados e controlos saudáveis comparáveis.",
        "Patrocínio ético e governação de consentimento, amostras e dados clínicos codificados.",
        "Capacidade laboratorial para isolamento de PBMC/CD8, controlo de RNA e RT-qPCR.",
        "Revisão conjunta de viabilidade, custos, procedimentos, autoria e partilha de dados antes do recrutamento.",
    ])
    doc.add_heading("Variáveis mínimas", level=1)
    paragraph(doc, "Idade, sexo, HLA-B27, tratamento, duração da doença, BASDAI, ASDAS, CRP/velocidade de sedimentação, atraso de processamento, viabilidade de PBMC, pureza CD8, qualidade de RNA e lote de RT-qPCR.")
    doc.add_heading("Critério de sucesso", level=1)
    paragraph(doc, "O piloto será bem-sucedido se produzir uma estimativa interpretável por participante, com controlo laboratorial fiável, independentemente de DDX24 ser inferior, semelhante ou superior. Um resultado negativo ou contraditório é cientificamente útil e poderá impedir o avanço de um mecanismo não sustentado.")
    doc.add_heading("Próxima conversa proposta", level=1)
    paragraph(doc, "Reunião de viabilidade de 30 minutos para confirmar acesso a participantes, circuito de amostras, percurso ético, custo esperado, propriedade de amostras/dados e responsáveis clínico, laboratorial e estatístico.")
    p = doc.add_paragraph()
    font(p.add_run("Investigadores: João Pais e Diana Koshman | Sem afiliação institucional"),
         size=9.5, bold=True)
    OUT.mkdir(parents=True, exist_ok=True)
    doc.save(SUMMARY)


def main():
    build_protocol()
    build_summary()
    print(PROTOCOL)
    print(SUMMARY)


if __name__ == "__main__":
    main()
