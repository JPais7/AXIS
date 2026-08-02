from __future__ import annotations

from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "data/publication/ddx24-study"
OUTPUT = OUT_DIR / "DDX24_formal_systematic_review_Joao_Pais_Diana_Koshman.docx"
FLOW = OUT_DIR / "figures/figure-3-interim-review-flow.png"
FOREST = OUT_DIR / "figures/figure-4-cd8-forest.png"

NAVY = RGBColor(31, 77, 120)
BLUE = RGBColor(46, 116, 181)
GRAY = RGBColor(90, 99, 110)
LIGHT = "F4F6F9"
GRID = "CAD3DE"
WHITE = RGBColor(255, 255, 255)


def set_font(run, size=11, bold=False, italic=False, color=None) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    total = sum(widths)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
        for cell, width in zip(row.cells, widths, strict=True):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)
    set_font(run, size=9, color=GRAY)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, NAVY, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True
    doc.styles["Heading 2"].paragraph_format.space_after = Pt(10)
    header = section.header.paragraphs[0]
    set_font(header.add_run("DDX24 systematic review and meta-analysis"), size=8.5, color=GRAY)
    add_page_number(section.footer.paragraphs[0])


def add_text(doc: Document, text: str, *, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.widow_control = True
    if bold_lead and text.startswith(bold_lead):
        set_font(p.add_run(bold_lead), bold=True)
        set_font(p.add_run(text[len(bold_lead):]))
    else:
        set_font(p.add_run(text))


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.194)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.208
        set_font(p.add_run(item))


def add_callout(doc: Document, title: str, text: str) -> None:
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(table, [9360])
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    set_font(p.add_run(title), size=10.5, bold=True, color=NAVY)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    set_font(p2.add_run(text), size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int],
              font_size=8.5) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, label in zip(table.rows[0].cells, headers, strict=True):
        set_cell_shading(cell, LIGHT)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(cell.paragraphs[0].add_run(label), size=9, bold=True, color=NAVY)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values, strict=True):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_font(cell.paragraphs[0].add_run(value), size=font_size)
    set_table_geometry(table, widths)


def make_figures() -> None:
    FLOW.parent.mkdir(parents=True, exist_ok=True)
    font_path = Path("C:/Windows/Fonts/arial.ttf")
    bold_path = Path("C:/Windows/Fonts/arialbd.ttf")
    font = ImageFont.truetype(str(font_path), 28)
    small = ImageFont.truetype(str(font_path), 21)
    bold = ImageFont.truetype(str(bold_path), 30)
    image = Image.new("RGB", (1500, 1180), "white")
    draw = ImageDraw.Draw(image)
    boxes = [
        (250, 35, 1250, 185, "47 Europe PMC records identified\nand screened"),
        (250, 250, 1250, 400, "38 records excluded at title/abstract\nwith recorded reasons"),
        (250, 465, 1250, 615, "9 full-text reports assessed"),
        (180, 680, 1320, 860, "2 reports excluded\n"
         "1 duplicate publication; 1 without donor replication"),
        (180, 925, 1320, 1115, "7 independent eligible designs\n"
         "3 processed public cohorts; 4 awaiting reusable matrices"),
    ]
    for left, top, right, bottom, label in boxes:
        draw.rounded_rectangle((left, top, right, bottom), radius=22,
                               fill="#F4F6F9", outline="#2E74B5", width=4)
        bbox = draw.multiline_textbbox((0, 0), label, font=font, spacing=10,
                                       align="center")
        x = (left + right - (bbox[2] - bbox[0])) / 2
        y = (top + bottom - (bbox[3] - bbox[1])) / 2
        draw.multiline_text((x, y), label, font=font, fill="#1F4D78",
                            spacing=10, align="center")
    for start, end in ((185, 250), (400, 465), (615, 680), (860, 925)):
        draw.line((750, start, 750, end - 18), fill="#2E74B5", width=5)
        draw.polygon([(735, end - 35), (765, end - 35), (750, end - 12)],
                     fill="#2E74B5")
    image.save(FLOW, dpi=(220, 220))

    studies = ["GSE194315", "GSE288581", "Random-effects"]
    effects = [-0.2299874100, -0.0992100595, -0.1480766703]
    lows = [-0.416360314, -0.233871408, -0.272079788]
    highs = [-0.043614506, 0.035451289, -0.024073552]
    image = Image.new("RGB", (1600, 760), "white")
    draw = ImageDraw.Draw(image)
    draw.text((90, 40), "DDX24 expression in compatible donor-level CD8 cohorts",
              font=bold, fill="#1F4D78")
    plot_left, plot_right = 480, 1120
    min_x, max_x = -0.52, 0.12
    def xpos(value: float) -> int:
        return int(plot_left + (value - min_x) / (max_x - min_x) * (plot_right - plot_left))
    zero_x = xpos(0)
    for y in range(120, 585, 18):
        draw.line((zero_x, y, zero_x, min(y + 9, 585)), fill="#7B8490", width=2)
    ys = [190, 340, 510]
    for idx, (name, eff, low, high, y) in enumerate(zip(studies, effects, lows, highs, ys)):
        color = "#1F4D78" if idx < 2 else "#C85A17"
        draw.text((95, y - 17), name, font=font, fill="#111111")
        draw.line((xpos(low), y, xpos(high), y), fill=color, width=5)
        draw.line((xpos(low), y - 12, xpos(low), y + 12), fill=color, width=4)
        draw.line((xpos(high), y - 12, xpos(high), y + 12), fill=color, width=4)
        cx = xpos(eff)
        if idx < 2:
            draw.rectangle((cx - 10, y - 10, cx + 10, y + 10), fill=color)
        else:
            draw.polygon([(cx, y - 15), (cx + 15, y), (cx, y + 15), (cx - 15, y)],
                         fill=color)
        draw.text((1160, y - 17), f"{eff:.3f} [{low:.3f}, {high:.3f}]",
                  font=small, fill="#111111")
    draw.line((plot_left, 610, plot_right, 610), fill="#111111", width=2)
    for value in (-0.5, -0.4, -0.3, -0.2, -0.1, 0):
        x = xpos(value)
        draw.line((x, 602, x, 618), fill="#111111", width=2)
        draw.text((x - 24, 625), f"{value:.1f}", font=small, fill="#111111")
    draw.text((520, 680), "Case-minus-control log2-CPM effect",
              font=small, fill="#111111")
    draw.text((95, 590), "Random-effects: I²=19.5%; tau²=0.00167; p=0.0193",
              font=small, fill="#5A636E")
    image.save(FOREST, dpi=(220, 220))


def add_cover(doc: Document) -> None:
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    set_font(p.add_run("SYSTEMATIC REVIEW AND META-ANALYSIS"), size=10,
             bold=True, color=BLUE)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(16)
    set_font(title.add_run(
        "DDX24 expression in peripheral CD8 T cells in ankylosing spondylitis"
    ), size=26, bold=True, color=NAVY)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(28)
    set_font(subtitle.add_run(
        "A targeted living systematic review and context-stratified meta-analysis "
        "of public human transcriptomic cohorts"
    ), size=14, italic=True, color=GRAY)
    authors = doc.add_paragraph()
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(authors.add_run("João Pais  •  Diana Koshman"), size=12, bold=True)
    aff = doc.add_paragraph()
    aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
    aff.paragraph_format.space_after = Pt(54)
    set_font(aff.add_run("No institutional affiliation"), size=10.5,
             italic=True, color=GRAY)
    status = doc.add_paragraph()
    status.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(status.add_run(
        "Formal manuscript draft • Evidence searched to 30 July 2026"
    ), size=9.5, color=GRAY)
    doc.add_page_break()


def add_front_matter(doc: Document) -> None:
    doc.add_heading("Abstract", level=1)
    for heading, text in (
        ("Background", "DDX24 regulates cytosolic RNA-mediated innate immune "
         "signalling, but its expression in ankylosing spondylitis has not been "
         "systematically synthesized at donor level."),
        ("Methods", "We conducted a targeted living review of public human "
         "case-control transcriptomic cohorts. The primary outcome was "
         "case-minus-control DDX24 expression in peripheral memory or effector "
         "CD8 T cells, with the human participant as the statistical unit. "
         "Compatible log2-CPM effects were synthesized by inverse-variance "
         "fixed-effect and DerSimonian-Laird random-effects models. Other blood "
         "platforms were synthesized only within compatible strata or described "
         "directionally. PubMed, Europe PMC and public repositories were searched "
         "to 30 July 2026. Risk of bias and certainty were assessed using "
         "prespecified domains."),
        ("Results", "Forty-seven Europe PMC records were screened, nine full texts "
         "were assessed and seven independent eligible designs were identified. "
         "Two independent "
         "cohorts with 47 participants were eligible for the primary quantitative "
         "synthesis. DDX24 was lower in cases in both cohorts. The pooled "
         "random-effects estimate was -0.148 log2-CPM (95% CI -0.272 to -0.024; "
         "p=0.0193; I²=19.5%). A broad-CD8 sensitivity analysis added GSE163314 "
         "(total 51 participants) and remained directionally consistent (-0.145; "
         "95% CI -0.249 to -0.041; p=0.0061). Three peripheral-blood microarray cohorts were "
         "directionally concordant but heterogeneous (pooled effect -0.267; "
         "95% CI -0.594 to 0.060; I²=86.7%). Two whole-blood RNA-sequencing "
         "cohorts were directionally mixed. Missing clinical covariates and the "
         "small number of CD8 cohorts lowered certainty."),
        ("Conclusions", "DDX24 shows a reproducible within-CD8 reduction across "
         "three independent public cohorts, although the third used a broader CD8 "
         "state and was retained as sensitivity evidence. Broader blood "
         "support. The evidence strengthens a falsifiable biological hypothesis "
         "but does not establish causality, druggability or therapeutic benefit."),
    ):
        doc.add_heading(heading, level=2)
        add_text(doc, text)
    keywords = doc.add_paragraph()
    keywords.paragraph_format.space_before = Pt(6)
    keywords.paragraph_format.space_after = Pt(10)
    set_font(keywords.add_run(
        "Keywords: ankylosing spondylitis; axial spondyloarthritis; DDX24; "
        "CD8 T cells; transcriptomics; systematic review; meta-analysis."
    ))
    doc.add_heading("Strengths and limitations of this review", level=1)
    add_bullets(doc, [
        "The participant, rather than the cell, was the statistical unit, avoiding cell-level pseudoreplication.",
        "The primary CD8 hypothesis and direction were frozen before incorporation of the next dataset.",
        "Assay-incompatible effects were not pooled into a misleading global estimate.",
        "The primary synthesis contains only two compatible memory/effector CD8 cohorts; the third broad-CD8 cohort has four participants.",
        "Age, sex, medication, HLA-B27 status and disease activity were not consistently available.",
        "Europe PMC screening and full-text assessment are complete, but Europe PMC overlaps substantially with PubMed and a lower-overlap Embase or Web of Science search remains pending.",
    ])
    add_callout(doc, "Interpretation boundary",
                "The valid claim is an association between ankylosing spondylitis "
                "and lower donor-level DDX24 expression in selected peripheral CD8 "
                "states. No result in this review proves that DDX24 causes disease "
                "or that changing DDX24 will benefit patients.")


def add_introduction(doc: Document) -> None:
    doc.add_heading("Introduction", level=1)
    add_text(doc, "Ankylosing spondylitis is an immune-mediated inflammatory "
             "disease within the axial spondyloarthritis spectrum. Its peripheral "
             "blood transcriptome is heterogeneous because cohorts differ in "
             "phenotype, treatment, disease activity, cell composition and assay "
             "technology [1].")
    add_text(doc, "DDX24 is an RNA helicase reported to restrain cytosolic "
             "RNA-mediated innate immune signalling [2]. It emerged from the AXIS "
             "workflow as a recurrent candidate, but a cross-study signal can be "
             "misleading if cells are treated as independent observations or if "
             "incompatible expression scales are pooled.")
    add_text(doc, "We therefore asked a narrow review question: in humans with "
             "ankylosing spondylitis, is donor-level DDX24 RNA expression in "
             "peripheral memory or effector CD8 T cells different from healthy "
             "controls? A secondary objective was to determine whether direction "
             "was supported in independent whole-blood contexts without estimating "
             "an invalid cross-platform global effect.")


def add_methods(doc: Document) -> None:
    doc.add_heading("Methods", level=1)
    doc.add_heading("Protocol, reporting and registration", level=2)
    add_text(doc, "The review followed an a priori project protocol and is reported "
             "against PRISMA 2020 and PRISMA-S principles [3,4]. The protocol was "
             "frozen within the AXIS project before the next eligible dataset. "
             "It has not yet been prospectively registered in PROSPERO; this is "
             "reported as a limitation rather than retrospectively labelled as "
             "prospective registration [5].")
    doc.add_heading("Eligibility criteria", level=2)
    add_text(doc, "We included human case-control RNA-expression studies with "
             "separable peripheral CD8 cells, at least two biological participants "
             "per group, donor-resolved labels and an available expression matrix "
             "or donor-level effect estimate. We excluded stimulated-only "
             "contrasts, studies without healthy controls, pooled libraries without "
             "participant resolution, non-RNA assays and duplicate repository records.")
    doc.add_heading("Information sources and search strategy", level=2)
    add_text(doc, "Searches covered PubMed, Europe PMC, NCBI GEO/SRA and "
             "BioStudies/ArrayExpress from inception to 30 July 2026. The PubMed "
             "concepts combined ankylosing spondylitis or axial "
             "spondyloarthritis with single-cell transcriptomics and CD8 or "
             "peripheral blood terms. The reproducible Europe PMC title/abstract "
             "query retrieved 47 records. Repository searches were followed by "
             "accession-level auditing and citation chasing. An Embase or Web of "
             "Science search requires institutional access and remains pending. "
             "Complete executable strings are reproduced in Appendix 1.")
    doc.add_heading("Selection process", level=2)
    add_text(doc, "All 47 Europe PMC records received a title/abstract decision "
             "and reason: 38 were excluded and nine proceeded to full text. "
             "Full-text assessment identified seven independent eligible designs, "
             "one duplicate publication and one study without donor replication. "
             "Diana Koshman confirmed all decisions with zero disagreements after "
             "reviewer-1 decisions were available; this is transparently classified "
             "as non-blinded confirmation rather than independent blinded screening.")
    doc.add_heading("Data collection and outcomes", level=2)
    add_text(doc, "For each cohort we extracted accession, assay, CD8 state, numbers "
             "of case and control donors, effect, standard error, p value, scale, "
             "data availability and eligibility decision. The primary outcome was "
             "case-minus-control donor-pseudobulk log2-CPM DDX24 expression in "
             "memory or effector CD8 cells. ADA was analysed identically as a "
             "contextual comparison target.")
    doc.add_heading("Risk of bias", level=2)
    add_text(doc, "Prespecified domains were participant selection, phenotype "
             "definition, treatment and clinical confounding, sample processing, "
             "donor-level statistical analysis, missing data and selective "
             "reporting. Judgements were low risk, some concerns or high risk, "
             "with reasons recorded. These review-specific domains reflect the "
             "observational omics design and are not presented as a validated "
             "replacement for a disease-specific risk-of-bias instrument.")
    doc.add_heading("Effect measures and synthesis", level=2)
    add_text(doc, "Within compatible CD8 strata, inverse-variance fixed-effect and "
             "DerSimonian-Laird random-effects models were estimated. Effects are "
             "case-minus-control log2-CPM differences. We report 95% confidence "
             "intervals, Cochran Q, I² and tau-squared. CD8 TEM and CD8 Naive "
             "effects from the same GSE194315 donors were not counted as independent "
             "studies. Peripheral-blood microarrays were pooled only on their "
             "normalized log-expression scale. Whole-blood FPKM and long-read TPM "
             "effects were reported directionally. No global cross-platform effect "
             "or p value was estimated. Random-effects interpretation followed "
             "standard meta-analytic principles [6].")
    doc.add_heading("Sensitivity and certainty", level=2)
    add_text(doc, "Sensitivity analyses replaced GSE194315 CD8 TEM with CD8 Naive, "
             "adjusted available processing-batch principal components and omitted "
             "each participant in turn. A further cell-state sensitivity added the "
             "author-annotated broad CD8_T population from GSE163314. Certainty was judged across risk of bias, "
             "inconsistency, indirectness, imprecision and dissemination bias, "
             "using GRADE concepts adapted to an exposure/association question [7]. "
             "Funnel plots and regression tests were prespecified only when at "
             "least ten compatible studies became available.")


def add_results(doc: Document) -> None:
    doc.add_heading("Results", level=1)
    doc.add_heading("Study selection", level=2)
    add_text(doc, "Europe PMC returned 47 records. Thirty-eight were excluded at "
             "title/abstract and nine full texts were assessed. Seven independent "
             "eligible designs were identified; one report duplicated PRJNA1168183 "
             "and HRA001027 lacked donor replication. Three eligible cohorts had "
             "public processed matrices. GSE194315 and GSE288581 comprised the "
             "primary memory/effector CD8 synthesis (14 cases and 33 controls); "
             "GSE163314 contributed a broad-CD8 sensitivity cohort (2 cases and "
             "2 controls).")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic = p.add_run().add_picture(str(FLOW), width=Inches(5.9))
    pic._inline.docPr.set("descr", "Interim flow of cohort identification, assessment, exclusions and inclusion.")
    pic._inline.docPr.set("title", "Interim review flow")
    cap = doc.add_paragraph()
    set_font(cap.add_run(
        "Figure 1. Europe PMC selection flow. Seven independent eligible designs "
        "were identified; availability and CD8-state compatibility determined "
        "their role in quantitative synthesis."
    ), size=9.5, italic=True, color=GRAY)

    doc.add_heading("Searches completed", level=2)
    add_table(doc,
              ["Source", "Date", "Scope", "Status"],
              [
                  ["PubMed", "30 Jul 2026", "Inception to date", "Primary pass executed"],
                  ["Europe PMC", "30 Jul 2026", "Inception to date", "47 records; screening complete"],
                  ["NCBI GEO/SRA", "30 Jul 2026", "Public repository", "Accessions audited"],
                  ["BioStudies/ArrayExpress", "30 Jul 2026", "Public repository", "Primary pass executed"],
                  ["Embase or Web of Science", "Pending", "Inception to date", "Institutional access required"],
              ],
              [1700, 1500, 2500, 3660], font_size=8.2)
    doc.add_heading("Candidate CD8 cohort registry", level=2)
    add_table(doc,
              ["Cohort", "Cases/controls", "CD8 scope", "Decision", "Reason"],
              [
                  ["GSE194315", "10/29", "CD8 TEM", "Include", "Donor-level matrix"],
                  ["GSE288581", "4/4", "CD45RO+ memory", "Include", "Donor-level matrix"],
                  ["GSE163314", "2/2", "Broad CD8_T", "Sensitivity", "Processed public matrix"],
                  ["PRJNA1168183", "14/3", "PBMC CD8 subsets", "Await matrix", "535.16 GB raw GEX only"],
                  ["PRJNA749866", "3/3", "PBMC CD8", "Raw only", "No processed matrix located"],
                  ["HRA001027", "1/1", "PBMC CD8", "Exclude", "One pooled library/group"],
                  ["GSE157595", "1/1", "PBMC CD8", "Mechanistic only", "scATAC, not RNA"],
                  ["GSE277791", "6/0", "PBMC CD8", "Exclude", "No healthy controls"],
                  ["NKG2C CD8 2025", "5/6", "PBMC CD8", "Contact authors", "Reusable matrix not located"],
                  ["BASSAA CITE-seq", "5/10", "PBMC CD8/TNK", "Contact authors", "Reusable matrix not located"],
              ],
              [1500, 1200, 1900, 1900, 2860], font_size=7.7)

    doc.add_heading("Primary CD8 meta-analysis", level=2)
    add_text(doc, "DDX24 was lower in cases in both eligible cohorts. GSE194315 "
             "contributed an effect of -0.230 log2-CPM (SE 0.0951), and GSE288581 "
             "contributed -0.099 (SE 0.0687). The random-effects pooled estimate "
             "was -0.148 (95% CI -0.272 to -0.024; p=0.0193). Heterogeneity was "
             "low but very imprecisely estimated with two studies (Q=1.243, "
             "p=0.265; I²=19.5%; tau²=0.00167).")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic = p.add_run().add_picture(str(FOREST), width=Inches(6.2))
    pic._inline.docPr.set("descr", "Forest plot of DDX24 effects in two donor-level CD8 cohorts and pooled random-effects estimate.")
    pic._inline.docPr.set("title", "CD8 DDX24 forest plot")
    cap = doc.add_paragraph()
    set_font(cap.add_run(
        "Figure 2. Donor-level DDX24 effects in compatible CD8 cohorts. Negative "
        "values indicate lower expression in cases. Confidence intervals are 95%."
    ), size=9.5, italic=True, color=GRAY)

    doc.add_heading("Sensitivity analyses", level=2)
    add_text(doc, "Replacing GSE194315 CD8 TEM with CD8 Naive preserved the "
             "lower-in-case direction but increased heterogeneity and uncertainty "
             "(random-effects estimate -0.262; 95% CI -0.597 to 0.074; p=0.126; "
             "I²=86.9%). Processing-batch adjustment in GSE194315 preserved lower "
             "DDX24 in CD8 TEM (-0.230; p=0.0211) and CD8 Naive (-0.442; p=0.00019). "
             "Every participant omission preserved the direction in the two "
             "GSE194315 CD8 states, and all eight donor omissions preserved direction "
             "in GSE288581. GSE163314 independently showed lower DDX24 in its "
             "author-annotated broad CD8_T population (-0.155; SE 0.169; p=0.478). "
             "Adding this cohort as a cell-state sensitivity produced a pooled "
             "effect of -0.145 log2-CPM (95% CI -0.249 to -0.041; p=0.0061; "
             "tau²=0) across three cohorts and 51 participants.")
    doc.add_heading("Support across blood contexts", level=2)
    add_text(doc, "Across all seven independent cohorts and 213 participants, "
             "DDX24 was lower in cases in six cohorts. Three peripheral-blood "
             "microarray cohorts were directionally concordant, but their pooled "
             "effect was imprecise and heterogeneous (-0.267; 95% CI -0.594 to "
             "0.060; p=0.109; I²=86.7%; tau²=0.0710). The two whole-blood "
             "RNA-sequencing cohorts were mixed. These contexts were not combined "
             "with the CD8 estimate.")
    add_table(doc,
              ["Context", "Cohorts", "Participants", "Direction", "Estimate or synthesis"],
              [
                  ["CD8 single-cell", "2", "47", "2/2 lower", "-0.148 [-0.272, -0.024]"],
                  ["Broad-CD8 sensitivity", "3", "51", "3/3 lower", "-0.145 [-0.249, -0.041]"],
                  ["Blood microarray", "3", "146", "3/3 lower", "-0.267 [-0.594, 0.060]"],
                  ["Whole-blood RNA-seq", "2", "20", "1/2 lower", "Directional only; mixed"],
              ],
              [2000, 1000, 1300, 1700, 3360], font_size=8.2)

    doc.add_heading("Risk of bias", level=2)
    add_text(doc, "All three processed cohorts used donor-level analysis and available "
             "case-control labels. Major concerns arose from unavailable clinical "
             "covariates; GSE288581 and GSE163314 also had very small samples. Processing and "
             "participant-influence checks reduced, but did not eliminate, bias concerns.")
    add_table(doc,
              ["Domain", "GSE194315", "GSE288581", "GSE163314", "Reason"],
              [
                  ["Participant selection", "Some concerns", "High", "High", "Convenience cohorts; n=8 and n=4 external cohorts"],
                  ["Phenotype definition", "Some concerns", "Some concerns", "Some concerns", "Repository phenotype; limited harmonisation"],
                  ["Clinical confounding", "High", "High", "High", "Age, sex, medication and activity incompletely available"],
                  ["Sample processing", "Some concerns", "Some concerns", "Some concerns", "Residual technical effects possible"],
                  ["Donor-level analysis", "Low", "Low", "Low", "Donors, not cells, were statistical units"],
                  ["Missing data", "Some concerns", "Some concerns", "Some concerns", "Clinical and processing metadata incomplete"],
                  ["Selective reporting", "Some concerns", "Some concerns", "Some concerns", "No prospective DDX24 plans"],
                  ["Overall", "High", "High", "High", "Confounding and small samples dominate"],
              ],
              [1700, 1200, 1200, 1200, 4060], font_size=7.2)

    doc.add_heading("Certainty of evidence", level=2)
    add_table(doc,
              ["Outcome", "Studies / participants", "Effect", "Certainty", "Main reasons"],
              [
                  ["DDX24 in memory/effector CD8", "2 / 47",
                   "-0.148 log2-CPM [-0.272, -0.024]", "Low",
                   "Risk of confounding; two studies; small external cohort"],
                  ["Broad-CD8 sensitivity", "3 / 51",
                   "-0.145 log2-CPM [-0.249, -0.041]", "Low",
                   "Indirect broad CD8 state; n=4 third cohort; confounding"],
                  ["DDX24 in blood microarrays", "3 / 146",
                   "-0.267 [-0.594, 0.060]", "Very low",
                   "High heterogeneity; imprecision; indirect bulk context"],
                  ["DDX24 in whole-blood RNA-seq", "2 / 20",
                   "Mixed direction", "Very low",
                   "Incompatible scales; imprecision; inconsistency"],
              ],
              [2100, 1500, 2200, 1000, 2560], font_size=7.5)


def add_discussion(doc: Document) -> None:
    doc.add_heading("Discussion", level=1)
    doc.add_heading("Principal findings", level=2)
    add_text(doc, "The principal finding is a modest but directionally consistent "
             "reduction of DDX24 in donor-level peripheral CD8 populations from "
             "people with ankylosing spondylitis. The association was present in "
             "two compatible memory/effector datasets, survived participant-influence analyses "
             "and was directionally replicated in a third independent broad-CD8 cohort, "
             "and was supported by direction in three independent blood microarray "
             "cohorts. The evidence was not universal: whole-blood RNA-sequencing "
             "cohorts disagreed, and CD8-state substitution increased heterogeneity.")
    doc.add_heading("Biological interpretation", level=2)
    add_text(doc, "DDX24 has been linked experimentally to regulation of cytosolic "
             "RNA sensing [2]. Reduced expression could mark altered RNA-processing "
             "or innate immune state, but the present data cannot distinguish a "
             "causal mechanism from consequences of inflammation, treatment or "
             "cell state. The broader direction across blood microarrays supports "
             "recurrence, while the mixed RNA-sequencing results caution against a "
             "universal whole-blood biomarker interpretation.")
    doc.add_heading("Comparison with other evidence", level=2)
    add_text(doc, "The included single-cell studies described altered immune states "
             "and CD8 biology in ankylosing spondylitis [8,9]. This review adds a "
             "gene-specific, donor-level cross-cohort synthesis. ADA did not show "
             "the same cross-context pattern: it was higher in CD8 cohorts, lower "
             "in microarrays and mixed in RNA-sequencing, demonstrating why a "
             "global cross-platform pool would be misleading.")
    doc.add_heading("Strengths", level=2)
    add_text(doc, "The design used participant-level pseudobulk effects, prespecified "
             "a primary CD8 context, froze the hypothesis before adding the next "
             "dataset, preserved exclusion reasons and checksums, separated assay "
             "contexts and performed batch and influence sensitivity analyses.")
    doc.add_heading("Limitations", level=2)
    add_text(doc, "Only two compatible memory/effector CD8 cohorts were eligible "
             "for the primary synthesis, one with eight participants. The third "
             "cohort contained only four participants and used a broader CD8_T state. "
             "Confidence intervals and heterogeneity estimates are therefore "
             "unstable. CD8 TEM and CD45RO-positive memory cells are related but "
             "not identical. Age, sex, medication, HLA-B27 status, disease duration "
             "and activity could not be harmonized. The bibliographic search is not "
             "yet complete in a lower-overlap subscription database, screening "
             "received non-blinded confirmation rather than independent blinded "
             "duplication, the protocol was not "
             "prospectively registered, and publication bias cannot be assessed "
             "with so few studies. Risk-of-bias judgements require independent "
             "confirmation.")
    doc.add_heading("Implications for research", level=2)
    add_text(doc, "The next decisive study is prospective validation of DDX24 in "
             "purified CD8 cells using RT-qPCR or targeted RNA sequencing, with "
             "matched controls and prespecified adjustment for age, sex, medication, "
             "HLA-B27 and disease activity. Reusable donor-level matrices should be "
             "requested for PRJNA1168183, PRJNA749866, the NKG2C CD8 cohort and "
             "the BASSAA CITE-seq cohort. "
             "Functional perturbation should follow only after expression replication "
             "and should include innate RNA-sensing, cell fitness and nonspecific "
             "RNA-processing safety endpoints.")
    doc.add_heading("Conclusions", level=1)
    add_text(doc, "Public donor-level transcriptomic evidence supports recurrent "
             "lower DDX24 expression within memory or effector CD8 T cells in "
             "ankylosing spondylitis, with consistent sensitivity evidence in a "
             "third broad-CD8 cohort. Certainty remains low because the primary "
             "synthesis contains only two compatible cohorts and clinical confounding is "
             "unresolved. The finding justifies independent laboratory falsification, "
             "not therapeutic claims.")


def add_declarations(doc: Document) -> None:
    doc.add_heading("Declarations", level=1)
    for heading, text in (
        ("Ethics approval", "This secondary analysis used publicly available, "
         "de-identified datasets. Any prospective sample collection will require "
         "local ethics approval and informed consent."),
        ("Consent for publication", "Not applicable."),
        ("Availability of data and materials", "Public accessions, derived tables, "
         "frozen inputs, checksums, analysis code and eligibility decisions are "
         "preserved in the AXIS project reproducibility manifest."),
        ("Competing interests", "The authors declare no competing interests."),
        ("Funding", "No external funding was declared."),
        ("Author contributions", "João Pais and Diana Koshman jointly conceived the "
         "review, interpreted the evidence and are responsible for final verification "
         "of the manuscript. Contributor roles must be confirmed by both authors "
         "before submission."),
        ("Use of automated methods", "The AXIS software supported reproducible data "
         "processing, evidence synthesis and document generation. The authors remain "
         "responsible for source verification, methodological decisions and all claims."),
    ):
        doc.add_heading(heading, level=2)
        add_text(doc, text)


def add_references(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("References", level=1)
    refs = [
        "Mauro D, et al. Ankylosing spondylitis: an autoimmune or autoinflammatory disease? Nat Rev Rheumatol. 2021;17:387-404. doi:10.1038/s41584-021-00625-y.",
        "Ma Z, et al. DDX24 negatively regulates cytosolic RNA-mediated innate immune signaling. PLoS Pathog. 2013;9. PMID:24204270.",
        "Page MJ, et al. The PRISMA 2020 statement: an updated guideline for reporting systematic reviews. BMJ. 2021;372:n71. doi:10.1136/bmj.n71.",
        "Rethlefsen ML, et al. PRISMA-S: an extension to the PRISMA Statement for reporting literature searches in systematic reviews. Syst Rev. 2021;10:39. doi:10.1186/s13643-020-01542-z.",
        "Centre for Reviews and Dissemination. PROSPERO: guidance for registering human studies. University of York.",
        "Higgins JPT, et al., editors. Cochrane Handbook for Systematic Reviews of Interventions. Chapter 10: Analysing data and undertaking meta-analyses. Version 6.5.",
        "GRADE Working Group. GRADE Book: assessing the certainty of evidence. Version 1.0, updated 2024-2026.",
        "Alber S, et al. Single cell transcriptome and surface epitope analysis of ankylosing spondylitis facilitates disease classification by machine learning. Front Immunol. 2022;13:838636. doi:10.3389/fimmu.2022.838636.",
        "Tang M, Qaiyum Z, Lim M, Inman RD. Single cell immune profiling in ankylosing spondylitis reveals resistance of CD8+ T cells to immune exhaustion. iScience. 2025;28:112715. doi:10.1016/j.isci.2025.112715.",
        "NCBI Gene Expression Omnibus. GSE194315: RNA and surface epitope sequencing of single cells involved in spondyloarthritis.",
        "NCBI Gene Expression Omnibus. GSE288581: Single Cell Immune Profiling in Ankylosing Spondylitis.",
        "Lefferts AR, et al. Circulating mature granzyme B+ T cells distinguish Crohn's disease-associated axial spondyloarthritis from axial spondyloarthritis and Crohn's disease. Arthritis Res Ther. 2021;23:147. doi:10.1186/s13075-021-02531-w.",
        "NCBI Gene Expression Omnibus. GSE163314: Comparing co-morbid Crohn's-spondyloarthritis to each underlying disease.",
        "Shamseer L, et al. Preferred reporting items for systematic review and meta-analysis protocols (PRISMA-P) 2015. BMJ. 2015;349:g7647. doi:10.1136/bmj.g7647.",
    ]
    for ref in refs:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.208
        set_font(p.add_run(ref), size=9.5)


def add_appendices(doc: Document) -> None:
    def search_heading(text: str) -> None:
        doc.add_heading(text, level=2)
        gap = doc.add_paragraph()
        gap.paragraph_format.space_after = Pt(2)

    doc.add_page_break()
    doc.add_heading("Appendix 1. Search strategies", level=1)
    search_heading("PubMed")
    add_text(doc, '(ankylosing spondylitis OR axial spondyloarthritis) AND '
             '(single-cell OR "single cell") AND ("RNA sequencing" OR transcriptom*) '
             'AND (CD8 OR "peripheral blood" OR PBMC)')
    search_heading("NCBI GEO/SRA")
    add_text(doc, 'ankylosing spondylitis AND "Homo sapiens" AND '
             '"single-cell RNA sequencing"')
    search_heading("BioStudies/ArrayExpress")
    add_text(doc, '("ankylosing spondylitis" OR "axial spondyloarthritis") AND '
             '("single-cell" OR transcriptomics)')
    search_heading("Europe PMC")
    add_text(doc, '(TITLE_ABS:"ankylosing spondylitis" OR '
             'TITLE_ABS:"axial spondyloarthritis") AND '
             '(TITLE_ABS:"single cell" OR TITLE_ABS:"single-cell") AND '
             '(TITLE_ABS:transcriptom* OR TITLE_ABS:"RNA sequencing")')
    search_heading("Pending Embase/Web of Science translation")
    add_text(doc, "Translate disease, single-cell transcriptomics, CD8 and blood "
             "concepts using database-specific controlled vocabulary; export exact "
             "query, date and result count before final submission.")
    doc.add_heading("Appendix 2. PRISMA completion actions", level=1)
    add_callout(doc, "Submission status",
                "This document has the formal structure and quantitative content of "
                "a systematic review and meta-analysis, but it should not be submitted "
                "as a completed PRISMA review until the pending actions below are closed.")
    add_table(doc,
              ["Required action", "Current status", "Evidence needed for completion"],
              [
                  ["Second open bibliographic source", "Complete", "Europe PMC: 47 records, exact query and export preserved"],
                  ["Lower-overlap database", "Pending", "Embase or Web of Science query and export"],
                  ["Deduplication counts", "Complete", "Machine-readable record library and duplicate log"],
                  ["Screening confirmation", "Non-blinded", "Repeat blinded screening if required by target journal"],
                  ["Full-text exclusion table", "Complete", "Nine reports assessed with reasons"],
                  ["Risk-of-bias duplication", "Pending", "Second reviewer judgements and agreement"],
                  ["Protocol registration", "Not prospective", "Register future living update; disclose timing"],
                  ["Processed third CD8 cohort", "Complete", "GSE163314 analyzed as broad-CD8 sensitivity"],
                  ["Additional eligible matrices", "Pending", "Author replies or PRJNA reprocessing"],
                  ["Reference verification", "Pending", "Human verification of every citation and DOI"],
              ],
              [3000, 1800, 4560], font_size=8.2)


def main() -> None:
    make_figures()
    doc = Document()
    configure_document(doc)
    doc.core_properties.title = (
        "DDX24 expression in peripheral CD8 T cells in ankylosing spondylitis"
    )
    doc.core_properties.author = "João Pais; Diana Koshman"
    doc.core_properties.subject = "Targeted living systematic review and meta-analysis"
    add_cover(doc)
    add_front_matter(doc)
    add_introduction(doc)
    add_methods(doc)
    add_results(doc)
    add_discussion(doc)
    add_declarations(doc)
    add_references(doc)
    add_appendices(doc)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
