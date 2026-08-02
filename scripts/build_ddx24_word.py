from __future__ import annotations

import csv
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "data/publication/ddx24-study/manuscript-draft.md"
FIGURES = ROOT / "data/publication/ddx24-study/figures"
STUDIES = ROOT / "data/publication/ddx24-study/study-characteristics.tsv"
OUTPUT = ROOT / "data/publication/ddx24-study/DDX24_article_Joao_Pais_Diana_Koshman.docx"

NAVY = RGBColor(31, 77, 120)
BLUE = RGBColor(46, 116, 181)
GRAY = RGBColor(90, 98, 108)
LIGHT = "E8EEF5"


def set_font(run, *, size: float, bold: bool = False, italic: bool = False,
             color: RGBColor | None = None) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instruction, separate, text, end):
        run._r.append(element)
    set_font(run, size=9, color=GRAY)


def shade(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_width(cell, width_twips: int) -> None:
    properties = cell._tc.get_or_add_tcPr()
    width = properties.find(qn("w:tcW"))
    if width is None:
        width = OxmlElement("w:tcW")
        properties.append(width)
    width.set(qn("w:w"), str(width_twips))
    width.set(qn("w:type"), "dxa")


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 12),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, NAVY, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_cover(doc: Document) -> None:
    for _ in range(5):
        doc.add_paragraph()
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(kicker.add_run("ORIGINAL RESEARCH ARTICLE"), size=10, bold=True,
             color=BLUE)
    kicker.paragraph_format.space_after = Pt(18)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(14)
    set_font(
        title.add_run(
            "Recurrent context-dependent reduction of DDX24 in "
            "ankylosing spondylitis"
        ),
        size=25,
        bold=True,
        color=NAVY,
    )

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(28)
    set_font(
        subtitle.add_run(
            "A reproducible multimodal integration of seven independent "
            "human transcriptomic cohorts"
        ),
        size=13,
        italic=True,
        color=GRAY,
    )

    authors = doc.add_paragraph()
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(authors.add_run("João Pais  •  Diana Koshman"), size=13, bold=True)
    authors.paragraph_format.space_after = Pt(5)
    affiliation = doc.add_paragraph()
    affiliation.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(
        affiliation.add_run("No institutional affiliation"),
        size=10.5,
        italic=True,
        color=GRAY,
    )
    affiliation.paragraph_format.space_after = Pt(68)

    status = doc.add_paragraph()
    status.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(
        status.add_run(
            "Computational secondary analysis • 7 cohorts • 213 participants"
        ),
        size=10,
        color=GRAY,
    )
    doc.add_page_break()


def add_paragraph_with_emphasis(doc: Document, text: str) -> None:
    citation_rules = (
        ("DDX24 emerged", " [1]"),
        ("GSE194315", " [2]"),
        ("GSE288581", " [3]"),
    )
    for marker, citation in citation_rules:
        if marker in text and citation not in text:
            text = f"{text}{citation}"
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.widow_control = True
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            set_font(paragraph.add_run(part[2:-2]), size=11, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            set_font(paragraph.add_run(part[1:-1]), size=10.5, italic=True)
        else:
            set_font(paragraph.add_run(part), size=11)


def add_manuscript_body(doc: Document) -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()[1:]
    paragraph_lines: list[str] = []
    inserted_results = False
    skip_section = False

    def flush() -> None:
        if paragraph_lines:
            add_paragraph_with_emphasis(doc, " ".join(paragraph_lines))
            paragraph_lines.clear()

    for line in lines:
        stripped = line.strip()
        if not stripped:
            flush()
            continue
        if stripped.startswith("## "):
            flush()
            heading = stripped[3:]
            skip_section = heading == "Provisional title and claims"
            if skip_section:
                continue
            doc.add_heading(heading, level=1)
            if heading == "Results" and not inserted_results:
                inserted_results = True
            continue
        if skip_section:
            continue
        if stripped.startswith("### "):
            flush()
            heading = stripped[4:]
            doc.add_heading(heading, level=2)
            continue
        if stripped.startswith("- "):
            flush()
            p = doc.add_paragraph(style="List Bullet")
            set_font(p.add_run(stripped[2:]), size=11)
            p.paragraph_format.space_after = Pt(4)
            continue
        if stripped.startswith("Allowed claim:") or stripped.startswith(
            "Prohibited claims:"
        ):
            continue
        paragraph_lines.append(stripped)
    flush()


def add_figures(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("Figures", level=1)
    captions = [
        (
            FIGURES / "figure-1-cohort-effects.png",
            "Figure 1. DDX24 case-minus-control effects in seven independent "
            "cohorts, separated into compatible assay contexts. Negative "
            "values indicate lower expression in cases. Effects are not pooled "
            "across panels because assay scales differ.",
        ),
        (
            FIGURES / "figure-2-context-concordance.png",
            "Figure 2. Directional concordance for DDX24 and ADA. Values show "
            "the fraction of independent cohorts within each context whose "
            "case-minus-control effect was below zero.",
        ),
    ]
    for image, caption in captions:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.keep_with_next = True
        picture = p.add_run().add_picture(str(image), width=Inches(6.35))
        picture._inline.docPr.set("descr", caption)
        picture._inline.docPr.set("title", caption.split(".", 1)[0])
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        cp.paragraph_format.space_after = Pt(14)
        set_font(cp.add_run(caption), size=9.5, italic=True, color=GRAY)


def add_study_table(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("Table 1. Included cohorts", level=1)
    with STUDIES.open(encoding="utf-8", newline="") as source:
        rows = list(csv.DictReader(source, delimiter="\t"))
    headers = ("Cohort", "Context", "Cases", "Controls", "Effect", "Direction")
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = (1350, 3000, 900, 900, 1100, 2110)
    for index, (cell, label, width) in enumerate(
        zip(table.rows[0].cells, headers, widths, strict=True)
    ):
        set_cell_width(cell, width)
        shade(cell, LIGHT)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(cell.paragraphs[0].add_run(label), size=9, bold=True, color=NAVY)
    for row in rows:
        values = (
            row["cohort"],
            row["context"].replace("_", " "),
            row["case_samples"],
            row["control_samples"],
            f"{float(row['effect']):.3f}",
            row["direction"].replace("_", " "),
        )
        data_row = table.add_row()
        data_row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
        cells = data_row.cells
        for cell, value, width in zip(cells, values, widths, strict=True):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_font(cell.paragraphs[0].add_run(value), size=8.5)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(4)
    note.paragraph_format.space_after = Pt(10)
    set_font(
        note.add_run(
            "Note: Effects are comparable only within their prespecified assay "
            "context. A global cross-platform effect was not estimated."
        ),
        size=9,
        italic=True,
        color=GRAY,
    )


def add_references(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("References", level=1)
    references = [
        "Ma Z, et al. DDX24 negatively regulates cytosolic RNA-mediated innate "
        "immune signaling. PLoS Pathogens. 2013;9. PMID: 24204270.",
        "Alber S, et al. Single cell transcriptome and surface epitope analysis "
        "of ankylosing spondylitis facilitates disease classification by "
        "machine learning. Front Immunol. 2022;13:838636. "
        "doi:10.3389/fimmu.2022.838636.",
        "Tang M, Qaiyum Z, Lim M, Inman RD. Single cell immune profiling in "
        "ankylosing spondylitis reveals resistance of CD8+ T cells to immune "
        "exhaustion. iScience. 2025;28:112715. "
        "doi:10.1016/j.isci.2025.112715.",
        "Mauro D, et al. Ankylosing spondylitis: an autoimmune or "
        "autoinflammatory disease? Nat Rev Rheumatol. 2021;17:387-404. "
        "doi:10.1038/s41584-021-00625-y.",
        "Page MJ, et al. The PRISMA 2020 statement: an updated guideline for "
        "reporting systematic reviews. BMJ. 2021;372:n71. "
        "doi:10.1136/bmj.n71.",
        "NCBI Gene Expression Omnibus. GSE194315: RNA and surface epitope "
        "sequencing of single cells involved in spondyloarthritis.",
        "NCBI Gene Expression Omnibus. GSE288581: Single Cell Immune Profiling "
        "in Ankylosing Spondylitis.",
    ]
    for index, reference in enumerate(references, start=1):
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.28)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.paragraph_format.space_after = Pt(6)
        set_font(p.add_run(reference), size=10)


def main() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_font(header.add_run("DDX24 in ankylosing spondylitis"), size=9, color=GRAY)
    add_page_number(section.footer.paragraphs[0])
    add_cover(doc)
    add_manuscript_body(doc)
    add_figures(doc)
    add_study_table(doc)
    add_references(doc)
    properties = doc.core_properties
    properties.title = (
        "Recurrent context-dependent reduction of DDX24 in ankylosing spondylitis"
    )
    properties.author = "João Pais; Diana Koshman"
    properties.subject = "Computational transcriptomic integration"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
