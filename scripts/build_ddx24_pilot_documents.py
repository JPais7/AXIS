from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "data/publication/ddx24-pilot"
PROTOCOL = OUT / "DDX24_pilot_protocol_Joao_Pais_Diana_Koshman.docx"
SUMMARY = OUT / "DDX24_pilot_collaborator_summary_Joao_Pais_Diana_Koshman.docx"

NAVY = RGBColor(31, 77, 120)
BLUE = RGBColor(46, 116, 181)
GRAY = RGBColor(90, 99, 110)
LIGHT = "F3F6FA"


def font(run, size=10.5, bold=False, italic=False, color=None):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def configure(doc: Document, title: str):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for name, size, color in (
        ("Heading 1", 15, BLUE),
        ("Heading 2", 12, NAVY),
        ("Heading 3", 10.5, NAVY),
    ):
        style = doc.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.keep_with_next = True
    doc.core_properties.title = title
    doc.core_properties.author = "João Pais; Diana Koshman"
    header = section.header.paragraphs[0]
    font(header.add_run("DDX24 laboratory validation pilot"), size=8.5, color=GRAY)


def paragraph(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.widow_control = True
    if bold_lead and text.startswith(bold_lead):
        font(p.add_run(bold_lead), bold=True)
        font(p.add_run(text[len(bold_lead):]))
    else:
        font(p.add_run(text))
    return p


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.paragraph_format.space_after = Pt(3)
        font(p.add_run(item))


def set_cell(cell, fill=None):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = OxmlElement("w:tcMar")
    for side in ("top", "start", "bottom", "end"):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), "100")
        node.set(qn("w:type"), "dxa")
        margins.append(node)
    tc_pr.append(margins)
    if fill:
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)


def table(doc, headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for cell, label in zip(t.rows[0].cells, headers, strict=True):
        set_cell(cell, LIGHT)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        font(cell.paragraphs[0].add_run(label), size=8.5, bold=True, color=NAVY)
    t.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for values in rows:
        cells = t.add_row().cells
        for cell, value in zip(cells, values, strict=True):
            set_cell(cell)
            font(cell.paragraphs[0].add_run(value), size=8.2)
    grid = t._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        node = OxmlElement("w:gridCol")
        node.set(qn("w:w"), str(width))
        grid.append(node)
    for row in t.rows:
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
        for cell, width in zip(row.cells, widths, strict=True):
            tcw = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tcw.set(qn("w:w"), str(width))
            tcw.set(qn("w:type"), "dxa")
    return t


def title_block(doc, kicker, title, subtitle):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    font(p.add_run(kicker.upper()), size=9, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    font(p.add_run(title), size=22, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    font(p.add_run(subtitle), size=11.5, italic=True, color=GRAY)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    font(p.add_run("Investigators: João Pais and Diana Koshman | No institutional affiliation"),
         size=9.5, color=GRAY)


def build_protocol():
    doc = Document()
    configure(doc, "Pilot protocol: independent validation of DDX24 in CD8 T cells")
    title_block(
        doc,
        "Pilot protocol | Version 1.0 | 30 July 2026",
        "Independent laboratory validation of DDX24 expression in peripheral CD8 T cells",
        "Exploratory case-control pilot in ankylosing spondylitis",
    )
    doc.add_heading("Protocol synopsis", level=1)
    table(doc, ["Element", "Predeclared specification"], [
        ["Research question", "Is donor-level DDX24 RNA expression lower in purified peripheral CD8 T cells from adults with ankylosing spondylitis than in healthy controls?"],
        ["Design", "Exploratory, cross-sectional, case-control laboratory pilot."],
        ["Target sample", "12 cases and 12 controls; recruit up to 15 per group to allow technical or eligibility losses. Minimum analyzable target: 10 per group."],
        ["Primary material", "Fresh peripheral blood; PBMC isolation followed by untouched CD8 T-cell enrichment or validated cell sorting."],
        ["Primary assay", "RT-qPCR for DDX24, normalized to a validated geometric mean of at least two stable reference genes."],
        ["Primary endpoint", "Donor-level normalized DDX24 expression (ΔCt; higher ΔCt means lower DDX24 expression)."],
        ["Primary comparison", "Case versus healthy control, with the participant as the statistical unit."],
        ["Status", "Protocol for ethics, laboratory feasibility and collaborator review; not authorization to recruit or collect samples."],
    ], [2300, 7060])

    doc.add_heading("1. Rationale", level=1)
    paragraph(doc, "AXIS identified lower donor-level DDX24 expression in two compatible memory/effector CD8 cohorts. A third independent broad-CD8 cohort showed the same direction. The primary two-cohort estimate was -0.148 log2-CPM (95% CI -0.272 to -0.024), while the three-cohort sensitivity estimate was -0.145 (95% CI -0.249 to -0.041). These observational results require prospective falsification in an independent population before causal or therapeutic interpretation.")

    doc.add_heading("2. Objectives and hypotheses", level=1)
    paragraph(doc, "Primary objective: ", bold_lead="Primary objective: ")
    paragraph(doc, "Estimate the difference in normalized DDX24 RNA expression in purified peripheral CD8 T cells between ankylosing spondylitis cases and healthy controls.")
    bullets(doc, [
        "Primary directional hypothesis: DDX24 expression is lower in cases.",
        "Secondary objective: estimate technical feasibility, RNA yield, assay failure and between-participant variability.",
        "Exploratory objectives: assess associations with disease activity, HLA-B27 and treatment without treating them as confirmatory.",
    ])

    doc.add_heading("3. Study population", level=1)
    doc.add_heading("Cases", level=2)
    bullets(doc, [
        "Adults aged 18 years or older.",
        "Clinician-confirmed radiographic axial spondyloarthritis/ankylosing spondylitis using documented accepted criteria.",
        "Able to provide informed consent and a peripheral-blood sample.",
        "Treatment, activity and disease duration documented rather than used for post hoc exclusion.",
    ])
    doc.add_heading("Healthy controls", level=2)
    bullets(doc, [
        "Adults without inflammatory rheumatic disease, inflammatory bowel disease, psoriasis or active autoimmune disease.",
        "Frequency-matched to cases by age band and sex where feasible.",
        "No acute infection or vaccination in the prespecified recent window used by the ethics-approved protocol.",
    ])
    doc.add_heading("Common exclusions", level=2)
    bullets(doc, [
        "Acute febrile illness or clinically significant current infection.",
        "Inability to provide informed consent.",
        "Insufficient blood volume, failed PBMC isolation or RNA failing predeclared quality criteria.",
        "Any additional safety exclusion required by the clinical collection site.",
    ])

    doc.add_heading("4. Sample size and interpretation", level=1)
    paragraph(doc, "This is a variance-estimation and feasibility pilot, not a definitive efficacy study. With 10-15 participants per group, precision will be limited and small effects may not be statistically significant. The target of 12 analyzable participants per group balances feasibility with estimation of donor variability. A definitive study should be powered using the observed pilot variance and a biologically meaningful effect, not the most favorable observed pilot effect.")

    doc.add_heading("5. Clinical and pre-analytical data", level=1)
    table(doc, ["Category", "Minimum variables"], [
        ["Participant", "Study ID, age, sex at birth; no direct identifiers in the analysis file."],
        ["Disease", "Diagnostic criteria, HLA-B27, disease duration, BASDAI, ASDAS, CRP and ESR where available."],
        ["Treatment", "Current NSAID, conventional DMARD, TNF inhibitor, IL-17 inhibitor, JAK inhibitor, corticosteroid and recent treatment change."],
        ["Collection", "Date/time, fasting status if controlled, processing delay, blood volume, operator and batch."],
        ["Laboratory QC", "PBMC yield/viability, CD8 purity, RNA concentration/integrity, reverse-transcription batch, qPCR plate and exclusion reason."],
    ], [2200, 7160])

    doc.add_heading("6. Laboratory workflow", level=1)
    table(doc, ["Step", "Minimum requirement"], [
        ["Blood collection", "Use one anticoagulant and a fixed collection-to-processing window for every group."],
        ["PBMC isolation", "Apply the same validated density-gradient or equivalent protocol; record delay, yield and viability."],
        ["CD8 enrichment", "Prefer untouched negative selection to reduce activation; alternatively use a prespecified FACS gate. Record purity."],
        ["RNA and cDNA", "Use the same extraction lot strategy and input policy; include RNA and reverse-transcription controls."],
        ["qPCR assay", "Validate specificity and 90-110% efficiency. Use technical replicates, no-template and no-RT controls."],
        ["Reference genes", "Evaluate a small candidate panel such as RPLP0, TBP, HPRT1 or PPIA and preselect at least two stable genes; do not assume GAPDH stability."],
        ["Batch control", "Balance cases and controls across isolation days, extraction batches and qPCR plates. Randomize well positions."],
        ["Blinding", "Use coded samples so laboratory personnel do not need case/control status during processing and quantification."],
    ], [1700, 7660])

    doc.add_heading("7. Quality-control rules", level=1)
    bullets(doc, [
        "Define PBMC viability, CD8 purity, RNA quality and replicate-variation thresholds before opening group labels.",
        "Repeat a measurement only for a documented technical failure, never because the biological result is unexpected.",
        "Resolve discordant technical replicates using a written rule and retain the original values in the audit trail.",
        "Freeze exclusions before statistical comparison and report every excluded participant with a non-identifying reason.",
    ])

    doc.add_heading("8. Statistical analysis plan", level=1)
    bullets(doc, [
        "Statistical unit: participant. Technical replicates are averaged within participant.",
        "Primary outcome: ΔCt using the geometric mean of validated reference genes; report the group difference with a 95% confidence interval.",
        "Primary comparison: two-sided Welch comparison or equivalent linear model; retain the predeclared lower-in-case directional interpretation without converting it to a one-sided significance test.",
        "Report 2^-ΔΔCt fold change as an interpretable secondary effect, with uncertainty.",
        "Use one parsimonious adjusted sensitivity model with group, age and sex only if data completeness and sample size permit. Treatment and activity analyses are exploratory.",
        "Display every participant value. Report missingness, assay failures and results with and without justified influential observations.",
        "Do not claim absence of an effect from a non-significant small pilot and do not infer causality or drug suitability.",
    ])

    doc.add_heading("9. Bias control and reproducibility", level=1)
    bullets(doc, [
        "Prospectively timestamp the protocol and analysis script before unblinding.",
        "Match collection and processing procedures between groups.",
        "Preserve anonymized raw Ct values, plate maps, QC decisions, code, software versions and file checksums.",
        "Separate confirmatory DDX24 analysis from exploratory genes and pathways.",
        "Have a second analyst verify exclusions, group labels and the primary result.",
    ])

    doc.add_heading("10. Ethics and governance", level=1)
    paragraph(doc, "No recruitment or blood collection may begin before approval by the responsible ethics committee and authorization from the clinical and laboratory institutions. Participants must provide informed consent covering blood collection, molecular analysis, coded clinical variables, data retention, secondary analysis and publication. The linkage key must remain at the clinical site and outside the research analysis dataset. Applicable data-protection and biological-sample rules must be confirmed locally.")

    doc.add_heading("11. Feasibility milestones", level=1)
    table(doc, ["Milestone", "Go criterion"], [
        ["Laboratory partnership", "Named clinical and laboratory leads accept the workflow and responsibilities."],
        ["Ethics package", "Protocol, consent, data sheet and data-management plan approved."],
        ["Technical rehearsal", "A small non-comparative test confirms CD8 purity, RNA yield and qPCR performance."],
        ["Pilot recruitment", "At least 10 analyzable participants per group with balanced processing."],
        ["Primary analysis", "Blinded data lock, predefined QC applied and participant-level effect reported."],
        ["Decision", "Use observed variance and feasibility to design replication; functional perturbation only after credible expression replication."],
    ], [2200, 7160])

    doc.add_heading("12. Roles requiring confirmation", level=1)
    table(doc, ["Role", "Required responsibility"], [
        ["Clinical principal investigator", "Diagnosis, eligibility, consent, safety and clinical governance."],
        ["Laboratory lead", "SOPs, biosafety, cell isolation, assay validation and QC."],
        ["Data custodian", "Pseudonymization, linkage-key protection and access control."],
        ["Statistical lead", "Frozen analysis plan, quality checks and reproducible reporting."],
        ["João Pais and Diana Koshman", "Research question, evidence synthesis, protocol coordination and interpretation; exact contributor roles to be agreed."],
    ], [2600, 6760])

    doc.add_heading("13. Pilot deliverables", level=1)
    bullets(doc, [
        "Ethics-approved protocol and consent materials.",
        "Laboratory SOP and qPCR validation record.",
        "Anonymized participant-level clinical and ΔCt dataset.",
        "Primary result with uncertainty, QC flow and complete exclusions.",
        "Decision memo for a larger replication or termination of the hypothesis.",
    ])
    OUT.mkdir(parents=True, exist_ok=True)
    doc.save(PROTOCOL)


def build_summary():
    doc = Document()
    configure(doc, "Collaborator summary: DDX24 validation pilot")
    title_block(
        doc,
        "Collaboration proposal",
        "DDX24 validation pilot in ankylosing spondylitis",
        "A small, falsifiable donor-level study in purified peripheral CD8 T cells",
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    font(p.add_run("Why this study"), size=12, bold=True, color=NAVY)
    paragraph(doc, "Across three independent public single-cell cohorts, donor-level DDX24 expression was lower in ankylosing spondylitis/axial spondyloarthritis. The primary compatible synthesis included 47 participants; a broad-CD8 sensitivity analysis included 51. Certainty remains low because cohorts are few, clinical covariates are incomplete and one cohort is very small. An independent laboratory pilot is therefore the next decisive test.")
    table(doc, ["Pilot element", "Proposal"], [
        ["Question", "Is DDX24 RNA expression lower in purified blood CD8 T cells from cases than controls?"],
        ["Participants", "Target 12 cases and 12 controls; recruit up to 15/group; minimum analyzable 10/group."],
        ["Assay", "PBMC isolation, untouched CD8 enrichment, RNA extraction and validated RT-qPCR."],
        ["Primary endpoint", "Participant-level normalized DDX24 ΔCt using at least two stable reference genes."],
        ["Bias control", "Balanced processing, coded samples, prespecified QC, participant-level statistics and complete audit trail."],
        ["Interpretation", "Exploratory replication only; no causal, diagnostic or therapeutic claim."],
    ], [2100, 7260])
    doc.add_heading("What we are seeking from a collaborator", level=1)
    bullets(doc, [
        "Clinical access to well-characterized ankylosing spondylitis participants and matched healthy controls.",
        "Ethics sponsorship and governance for consent, samples and coded clinical data.",
        "Laboratory capability for PBMC/CD8 isolation, RNA QC and RT-qPCR.",
        "Joint review of feasibility, costs, SOPs, authorship and data-sharing before recruitment.",
    ])
    doc.add_heading("Minimum variables", level=1)
    paragraph(doc, "Age, sex, HLA-B27, treatment, disease duration, BASDAI, ASDAS, CRP/ESR, processing delay, PBMC viability, CD8 purity, RNA QC and qPCR batch.")
    doc.add_heading("Success criterion", level=1)
    paragraph(doc, "The pilot succeeds if it produces an interpretable participant-level estimate with reliable laboratory QC, regardless of whether DDX24 is lower, unchanged or higher. A negative or contradictory result is scientifically valuable and may stop an unsupported mechanism from advancing.")
    doc.add_heading("Proposed next conversation", level=1)
    paragraph(doc, "A 30-minute feasibility meeting to confirm participant access, sample workflow, ethics route, expected cost, ownership of samples/data and named clinical, laboratory and statistical leads.")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    font(p.add_run("Investigators: João Pais and Diana Koshman | No institutional affiliation"),
         size=9.5, bold=True, color=NAVY)
    OUT.mkdir(parents=True, exist_ok=True)
    doc.save(SUMMARY)


def main():
    build_protocol()
    build_summary()
    print(PROTOCOL)
    print(SUMMARY)


if __name__ == "__main__":
    main()
