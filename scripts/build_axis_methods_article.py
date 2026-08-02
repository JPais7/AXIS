# ruff: noqa: E501

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "data" / "publication" / "axis-methods"
DOCX = OUT / "AXIS_methods_manuscript_Joao_Pais_Diana_Koshman.docx"
VALIDATION_FIGURE = (
    ROOT
    / "benchmarks/comparison/multi-cohort-validation/axis-limma-validation.png"
)

NAVY = "17365D"
BLUE = "2F75B5"
PALE = "EAF2F8"
GRAY = "666666"
WHITE = "FFFFFF"


def font(run, size=10.5, bold=False, italic=False, color="000000"):
    run.font.name = "Aptos"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Aptos")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Aptos")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            margins = OxmlElement("w:tcMar")
            for side in ("top", "left", "bottom", "right"):
                node = OxmlElement(f"w:{side}")
                node.set(qn("w:w"), "110")
                node.set(qn("w:type"), "dxa")
                margins.append(node)
            tc_pr.append(margins)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_body(doc, text, bold_start=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.12
    if bold_start and text.startswith(bold_start):
        font(p.add_run(bold_start), bold=True)
        font(p.add_run(text[len(bold_start) :]))
    else:
        font(p.add_run(text))
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        font(p.add_run(item), size=10.2)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, value in enumerate(headers):
        cell = table.rows[0].cells[index]
        shade(cell, NAVY)
        p = cell.paragraphs[0]
        font(p.add_run(value), size=9.2, bold=True, color=WHITE)
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(row):
            if row_index % 2 == 1:
                shade(cells[index], "F5F7FA")
            p = cells[index].paragraphs[0]
            font(p.add_run(value), size=9.0)
    set_table_geometry(table, widths)
    repeat_header(table.rows[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def configure(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    for name, size, color in (("Heading 1", 16, NAVY), ("Heading 2", 12.5, BLUE), ("Heading 3", 11, NAVY)):
        style = styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12 if name == "Heading 1" else 8)
        style.paragraph_format.space_after = Pt(4)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font(header.add_run("AXIS | Software and methods manuscript | Draft 2.0"), size=8.5, color=GRAY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)


def build():
    OUT.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(72)
    p.paragraph_format.space_after = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(p.add_run("AXIS"), size=30, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    font(p.add_run("An auditable, participant-aware workflow for cross-study molecular evidence synthesis and therapeutic hypothesis generation"), size=15, color=BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(p.add_run("João Pais and Diana Koshman"), size=12, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(p.add_run("Independent researchers; no institutional affiliation"), size=10, italic=True, color=GRAY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    font(p.add_run("Software and methods manuscript - pre-submission draft"), size=10.5, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(p.add_run("Draft 2.0 | 2 August 2026"), size=9.5, color=GRAY)

    doc.add_page_break()
    add_heading(doc, "Abstract")
    add_body(doc, "Public molecular datasets can support therapeutic hypothesis generation, but reuse is hindered by inconsistent sample metadata, incompatible tissues and assays, pseudoreplication, and weak provenance. AXIS is an open-source Python platform that organizes study discovery, eligibility review, expression analysis, cross-cohort synthesis, single-cell pseudobulk validation, target evidence integration, and frozen reproduction as one guarded workflow. It distinguishes discovery, independent validation, sensitivity analysis, mechanistic evidence, and treatment response; preserves participant-level independence; and records checksums and method decisions for audit. In an axial-spondyloarthritis case study, AXIS synthesized DDX24 expression in two compatible CD8 cohorts (14 cases and 33 controls) and retained a broader third CD8 cohort as sensitivity evidence. Technical validation against native limma covered four GEO cohorts, 163 samples, three microarray platforms and an additional covariate-adjusted analysis. Effect rankings, adjusted-p-value rankings and leading probe lists were concordant. The current suite contains 146 automated tests and passes on Linux, macOS and Windows. AXIS is intended as research software for generating falsifiable priorities, not as an autonomous drug-discovery or clinical decision system.")
    add_body(doc, "Keywords: axial spondyloarthritis; transcriptomics; single-cell RNA sequencing; pseudobulk; meta-analysis; target prioritization; reproducible research")

    add_heading(doc, "1. Summary")
    add_body(doc, "AXIS converts heterogeneous public molecular evidence into traceable research hypotheses. Its central design principle is that more records do not automatically mean more evidence: studies enter a synthesis only when their disease definition, biological material, assay, comparison, and participant structure support the declared question. In single-cell analyses, cells are aggregated within participants so that the biological participant, rather than each cell, remains the statistical unit. The platform then keeps compatible primary evidence separate from broader sensitivity evidence.")

    add_heading(doc, "2. Statement of need")
    add_body(doc, "Repositories such as the NCBI Gene Expression Omnibus provide access to high-throughput functional genomics studies [1]. However, the metadata needed for secondary disease research are frequently encoded in free text, and apparently similar studies may differ in tissue, cell state, treatment exposure, platform, or unit of replication. A naive pipeline can therefore inflate sample size, combine incompatible effects, or rediscover the same participants under multiple accessions.")
    add_body(doc, "AXIS addresses this gap by making eligibility, provenance, and allowed scientific use first-class objects. It is designed for investigators who need to move from a broad repository search to a small set of defensible, testable claims without concealing uncertainty behind a single opaque score.")

    add_heading(doc, "3. Software design")
    add_table(doc, ["Stage", "Purpose", "Guardrail"], [
        ("Discovery", "Search and deduplicate GEO, BioStudies/ArrayExpress and SRA records.", "A repository record is not counted as an eligible cohort."),
        ("Eligibility", "Review disease, tissue, assay, design, treatment and participant independence.", "Approvals are role-specific and invalidated when analyzed inputs change."),
        ("Analysis", "Run QC, bulk differential analysis, or participant-level single-cell pseudobulk.", "Outliers are flagged, not silently removed; cells are not independent replicates."),
        ("Synthesis", "Assess recurrence, directional concordance, validation and random-effects summaries.", "Incompatible scales and cell definitions remain stratified."),
        ("Translation", "Join genetic, mechanistic, tractability, safety, structure and drug evidence.", "Expression direction is not substituted for causal therapeutic direction."),
        ("Reproduction", "Verify frozen inputs, lockfile, numerical outputs and claim constraints offline.", "Checksum or scientific-guardrail failures stop the run."),
    ], [1450, 3880, 4030])

    add_heading(doc, "3.1 Evidence and provenance model", 2)
    add_body(doc, "AXIS stores scientific entities, contextual claims, source assertions, calculated evidence, researcher hypotheses, and immutable provenance separately. A local DuckDB evidence store supports versioned migrations and append-only hypothesis revision. Download manifests record source locations, retrieval times, sizes, and SHA-256 checksums. Eligibility decisions are bound to checksums of the analyzed results, preventing an approval from silently surviving reanalysis.")

    add_heading(doc, "3.2 Statistical safeguards", 2)
    add_body(doc, "For microarrays, AXIS supports probe-to-gene mapping, multiple-testing correction, a two-group moderated model, and declared covariate designs. Cross-study recurrence requires same-direction support and does not pool incompatible raw effect scales. For single-cell RNA sequencing, counts are aggregated by participant and cell type before comparison. This follows evidence that methods accounting for the dependence of cells within individuals, including pseudobulk approaches, provide more reliable differential-expression inference [2]. Random-effects summaries are reserved for contrasts with sufficiently compatible definitions.")

    add_heading(doc, "3.3 Therapeutic evidence", 2)
    add_body(doc, "Candidate genes can be linked to Open Targets evidence covering disease association, human genetics, tractability, clinical precedent, safety and target-prioritization properties [3,4]. These dimensions remain visible rather than being collapsed into a proprietary score. Protein structure and pharmacology are downstream filters: neither an AlphaFold model nor an existing ligand can rescue a target that lacks replicated biological evidence or a defensible causal direction.")

    add_heading(doc, "4. Case study: DDX24 in CD8 T cells")
    add_body(doc, "The current demonstration asks whether DDX24 expression differs in CD8 T cells from people with axial or ankylosing spondylitis and healthy controls. Two compatible participant-level cohorts form the primary synthesis: 14 cases and 33 controls. The random-effects estimate is -0.148 (95% confidence interval -0.272 to -0.024; p=0.019). A third, broader CD8 cohort is sensitivity-only because its cellular definition is less specific. With that cohort, the synthesis contains 51 participants and retains the same direction (estimate -0.145; 95% confidence interval -0.249 to -0.041; p=0.006). These results support an association, not causality, target validity, or treatment benefit.")
    add_table(doc, ["Evidence role", "Cohorts", "Participants", "Interpretation"], [
        ("Primary CD8 synthesis", "2", "47", "Compatible CD8 definitions; associative evidence."),
        ("Broad-CD8 sensitivity", "3", "51", "Directional robustness under a broader cellular definition."),
        ("Excluded pooled dataset", "0 added", "Not estimable", "One pooled library per group cannot establish participant-level replication."),
    ], [2050, 1100, 1400, 4810])

    add_heading(doc, "5. Verification and reproducibility")
    add_body(doc, "On 2 August 2026, the complete suite passed 146 of 146 tests. Ruff and strict mypy checks also passed. GitHub Actions reproduced the checks with Python 3.12 on Linux, macOS and Windows. A packaged synthetic demonstration runs without distributing participant-level source data. The DDX24 offline reproduction verifies three frozen participant-level inputs and the Poetry lockfile by SHA-256, recomputes the primary and sensitivity summaries, compares numerical values within tight tolerances, and enforces 25 scientific and computational checks.")

    add_heading(doc, "5.1 Native-limma technical validation", 2)
    add_body(doc, "AXIS differential-expression results were compared with native limma 3.68.4 under R 4.6.1 in four public GEO cohorts: GSE18781/GPL570, GSE25101/GPL6947, GSE73754/GPL10558 and GSE11886/GPL570. Together, the comparisons covered 163 samples and three platform contexts. Across all four unadjusted contrasts, every probe-level effect direction agreed, effect Spearman correlations were approximately 1.0, adjusted-p-value Spearman correlations were 1.0, and the top-100 and top-500 probe sets overlapped completely. Maximum absolute effect differences were below 5.2 x 10^-12.")
    add_body(doc, "A separate GSE73754 comparison used the declared covariate-adjusted design containing group, sex, age and numeric batch. Across 47,323 shared probes, effect and adjusted-p-value rankings were concordant, all directions agreed, the top-100 and top-500 sets overlapped completely, and the maximum absolute effect difference was 4.52 x 10^-12. An offline synthetic regression test freezes a native-limma reference so that routine continuous integration can detect future changes in coefficients, directions and rankings without requiring R.")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    picture = run.add_picture(str(VALIDATION_FIGURE), width=Inches(6.35))
    picture._inline.docPr.set(
        "descr",
        "Five-cohort comparison showing near-perfect agreement between AXIS "
        "and native limma for effect rankings, leading probes and numerical "
        "estimates.",
    )
    caption = doc.add_paragraph()
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_together = True
    font(caption.add_run("Figure 1. "), size=9, bold=True)
    font(caption.add_run("Technical agreement between AXIS and native limma. Four unadjusted GEO contrasts and one covariate-adjusted contrast show near-perfect effect correlation, complete leading-list overlap and numerical differences on the order of 10^-12. This validates implementation for the tested designs; it does not establish biological truth or clinical validity."), size=9)

    add_body(doc, "These results demonstrate cross-platform re-executability, claim binding and numerical agreement for the tested microarray designs. They do not demonstrate that AXIS is biologically correct in every setting, superior to limma, or equivalent for paired designs, RNA sequencing, single-cell models or untested covariate structures.")

    add_heading(doc, "6. Research impact and intended use")
    add_body(doc, "AXIS has already supported a structured axial-spondyloarthritis case study and the preparation of a prospective DDX24 laboratory validation protocol. Its intended impact is methodological: to make the path from public data to experiment selection more conservative, inspectable, and reproducible. The platform can be adapted through a research manifest defining disease, population, tissue, cell type, comparison, treatments, and exclusion criteria. It should be used to prioritize experiments and evidence gaps, not to recommend treatment or claim clinical efficacy.")

    add_heading(doc, "7. Limitations")
    add_bullets(doc, [
        "The current demonstration is disease-specific and based on few compatible cohorts.",
        "The native-limma comparison covers four public microarray cohorts; broader platform, disease and design coverage remains necessary.",
        "The offline regression fixture is intentionally small and protects rankings and coefficients more strongly than exact empirical-Bayes statistics.",
        "Benchmarking against additional end-to-end tools and formal runtime/memory profiling remain incomplete.",
        "Human review is still required for sample interpretation, eligibility, confounders and biological meaning.",
        "RCT evidence, pharmacovigilance and clinical outcomes are not yet implemented as a dedicated evidence layer.",
        "The public development history is recent and therefore does not yet demonstrate long-term community adoption or maintenance.",
    ])

    add_heading(doc, "8. Availability")
    availability = add_body(doc, "AXIS is implemented in Python 3.12 and exposes a command-line interface. Source code is publicly available at https://github.com/JPais7/AXIS under the Apache License 2.0. The archived Zenodo record is available at https://doi.org/10.5281/zenodo.21760202. The repository includes locked dependencies, automated cross-platform tests, a packaged synthetic demonstration, frozen reproduction manifests, validation scripts, checksums and auditable tabular/JSON outputs. Raw public data that can be retrieved reproducibly and frozen participant-level data that should not be redistributed are excluded from the software archive.")
    availability.paragraph_format.keep_together = True

    add_heading(doc, "9. AI usage disclosure")
    add_body(doc, "Generative AI was used interactively to assist software implementation, documentation, test design and manuscript drafting. Scientific claims were constrained by explicit programmatic checks, source artifacts and author review. AI output was not treated as evidence, and the authors remain responsible for verification, interpretation and the final submitted text.")

    add_heading(doc, "10. Author contributions")
    add_body(doc, "João Pais: conceptualization, investigation, software, validation, project administration, visualization, and writing - original draft and review. Diana Koshman: investigation, validation planning, and writing - review and editing. Software engineering and manuscript drafting included AI-assisted work under author supervision. Both authors must approve the submitted version.")

    add_heading(doc, "11. Funding, competing interests and data ethics")
    add_body(doc, "Funding: none declared. Competing interests: none declared. AXIS reuses de-identified public datasets and does not itself recruit participants. The ethical and consent conditions of each source study remain applicable. These statements must be reconfirmed before submission.")

    add_heading(doc, "References")
    refs = [
        "1. Barrett T, et al. NCBI GEO: archive for high-throughput functional genomic data. Nucleic Acids Research. 2009;37(Database issue):D885-D890. doi:10.1093/nar/gkn764.",
        "2. Murphy AE, Skene NG. A balanced measure shows superior performance of pseudobulk methods in single-cell RNA-sequencing analysis. Nature Communications. 2022;13:7851. doi:10.1038/s41467-022-35519-4.",
        "3. Ochoa D, et al. Open Targets Platform: supporting systematic drug-target identification and prioritisation. Nucleic Acids Research. 2021;49(D1):D1302-D1310. doi:10.1093/nar/gkaa1027.",
        "4. Open Targets Consortium. Open Targets Platform: facilitating therapeutic hypotheses building in drug discovery. Nucleic Acids Research. 2025;53(D1):D1467-D1475. doi:10.1093/nar/gkae1072.",
        "5. Alber S, et al. Single cell transcriptome and surface epitope analysis of ankylosing spondylitis facilitates disease classification by machine learning. Frontiers in Immunology. 2022;13:838636. doi:10.3389/fimmu.2022.838636.",
        "6. Tang M, Qaiyum Z, Lim M, Inman RD. Single cell immune profiling in ankylosing spondylitis reveals resistance of CD8+ T cells to immune exhaustion. iScience. 2025;28:112715. doi:10.1016/j.isci.2025.112715.",
        "7. Ritchie ME, et al. limma powers differential expression analyses for RNA-sequencing and microarray studies. Nucleic Acids Research. 2015;43(7):e47. doi:10.1093/nar/gkv007.",
        "8. Pais J, Koshman D. AXIS: AI for Axial Spondyloarthritis Insights & Solutions [software]. Zenodo. 2026. doi:10.5281/zenodo.21760202.",
    ]
    for ref in refs:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.paragraph_format.space_after = Pt(4)
        font(p.add_run(ref.split(". ", 1)[1]), size=9.2)

    props = doc.core_properties
    props.title = "AXIS methods manuscript"
    props.author = "João Pais; Diana Koshman"
    props.subject = "Research software and methods"
    props.keywords = "AXIS, transcriptomics, reproducibility, pseudobulk, target prioritization"
    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    build()
